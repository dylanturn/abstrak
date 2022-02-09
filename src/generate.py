import sys
import util
import jmespath
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from resume_style import margins, AbstrakStyle


#########################
###   FIRST HEADER    ###
#########################
def build_header(document, header_data):
  default_section = document.sections[0]
  default_section.different_first_page_header_footer = True
  first_page_header = default_section.header.paragraphs[0]
  header= document.sections[0].first_page_header.paragraphs[0]

  default_section.top_margin = margins.get("top")
  default_section.right_margin = margins.get("right")
  default_section.bottom_margin = margins.get("bottom")
  default_section.left_margin = margins.get("left")

  first,last = header_data["title"].split()
  city = header_data["location"]["city"]
  state = header_data["location"]["state"]
  zip = header_data["location"]["zip"]

  #####
  # FIRST PAGE HEADER
  #####
  first_page_header.add_run(f"{first} ", style="Abstrak Title")
  first_page_header.add_run(f"{last}\n", style="Abstrak Title Accent")

  #####
  # HEADER TITLE
  #####
  header.add_run(f"{first} ", style="Abstrak Title")
  header.add_run(f"{last}\n", style="Abstrak Title Accent")

  #####
  # HEADER LOCATION
  #####
  header.add_run(city, style="Abstrak Subtitle")
  header.add_run(" • ", style="Abstrak Subtitle Seperator")
  header.add_run(state, style="Abstrak Subtitle")
  header.add_run(" • ", style="Abstrak Subtitle Seperator")
  header.add_run(f"{zip}\n", style="Abstrak Subtitle")
  
  #####
  # CONTACT INFO
  #####
  build_contacts(header, header_data["contacts"])
  return header


####################
###    FOOTER    ###
####################
def build_footer(document, footer_data):
  default_section = document.sections[0]
  
  footer_paragraph = default_section.footer.paragraphs[0]
  footer_paragraph.add_run("\n")
  build_contacts(footer_paragraph, footer_data["contacts"])
  footer_paragraph.add_run("\t\t")
  util.add_page_number(footer_paragraph.add_run(), style=document.styles["Normal Accent Emphasis"])
  footer_paragraph.add_run("|")
  util.add_page_number(footer_paragraph.add_run(), "NUMPAGES")
  footer_tab_stops = footer_paragraph.paragraph_format.tab_stops
  footer_tab_stops.add_tab_stop(Inches(0), WD_TAB_ALIGNMENT.LEFT)
  footer_tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)

  
  footer_paragraph= document.sections[0].first_page_footer.paragraphs[0]
  footer_paragraph.add_run("\n")
  build_contacts(footer_paragraph, footer_data["contacts"])
  footer_paragraph.add_run("\t\t")
  util.add_page_number(footer_paragraph.add_run(), style=document.styles["Normal Accent Emphasis"])
  footer_paragraph.add_run("|")
  util.add_page_number(footer_paragraph.add_run(), "NUMPAGES")
  footer_tab_stops = footer_paragraph.paragraph_format.tab_stops
  footer_tab_stops.add_tab_stop(Inches(0), WD_TAB_ALIGNMENT.LEFT)
  footer_tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)


# Builds the contact methods in a consistent way for the header and footer.
def build_contacts(paragraph, contacts):
  paragraph.add_run(f"{contacts[0]['name']} ", style="Abstrak Subtitle Emphasis")
  paragraph.add_run(contacts[0]["value"], style="Abstrak Subtitle")
  paragraph.add_run(" • ", style="Abstrak Subtitle Seperator")

  for idx in range(1, len(contacts)-1):
    paragraph.add_run(f"{contacts[idx]['name']} ", style="Abstrak Subtitle Emphasis")
    paragraph.add_run(f"{contacts[idx]['value']}", style="Abstrak Subtitle")
    paragraph.add_run(" • ", style="Abstrak Subtitle Seperator")
    
  paragraph.add_run(f"{contacts[-1]['name']} ", style="Abstrak Subtitle Emphasis")
  paragraph.add_run(f"{contacts[-1]['value']}", style="Abstrak Subtitle")


#####################
###    SUMMARY    ###
#####################
def build_summary(document, summary_data):
  document.add_paragraph("summary", style="Heading 1")
  document.add_paragraph(summary_data).add_run()


########################
###    HIGHLIGHTS    ###
########################
def build_highlights(document, highlights_data):

  # Skillset
  document.add_paragraph(highlights_data["skillset"]["title"], style="Heading 2")
  for skill in highlights_data["skillset"]["skills"]:
    document.add_paragraph(
      skill,
      style='List Bullet'
    )

  # Personal Projects
  document.add_paragraph(highlights_data["personal_projects"]["title"], style="Heading 2")
  for project in highlights_data["personal_projects"]["projects"]:
    project_paragraph = document.add_paragraph(
      style='List Bullet'
    )
    util.add_hyperlink(project_paragraph, project['name'], project['url'])
    project_paragraph.add_run(f" - {project['description']}")


###################
###    ROLES    ###
###################
def build_roles(document, role_data):
  
  # Try build the volunteer roles
  document.add_paragraph("volunteer work", style="Heading 1")
  volunteer_role_expression = jmespath.compile("[?type=='volunteer']")
  volunteer_roles = volunteer_role_expression.search(role_data)
  for role in volunteer_roles:
    build_role_positions(document, role)

  # Try build the professional roles
  document.add_paragraph("experience", style="Heading 1")
  professional_role_expression = jmespath.compile("[?type=='professional']")
  professional_roles = professional_role_expression.search(role_data)
  for role in professional_roles:
    build_role_positions(document, role)


def build_role_positions(document, role_data):
  # Get the role org and location information
  org_name = role_data["organization"]["name"]
  org_city = role_data["organization"]["location"]["city"]
  org_state = role_data["organization"]["location"]["state"]
  
  # Get the information for each position in each role
  for position in role_data["positions"]:
    pos_title = position["title"]
    pos_start = f"{position['dates']['start']['month']}/{position['dates']['start']['year']}"
    if (not position['dates']['end'].get("month")) or (not position['dates']['end'].get("year")):
      pos_end = "Present"
    else:
      pos_end = f"{position['dates']['end']['month']}/{position['dates']['end']['year']}"
    if role_data["organization"]["location"]["type"] == "remote":
      pos_location = "Remote"
    else:
      pos_location = f"{org_city}, {org_state}"

    # Position Title
    document.add_paragraph(pos_title, style="Abstrak Position Title")

    # Position Company, State, Dates
    position_paragraph = document.add_paragraph()
    position_paragraph.add_run(org_name, style="Abstrak Position Subtitle")
    position_paragraph.add_run(f", {pos_location}", style="Abstrak Position Subtitle")
    position_paragraph.add_run(" • ", style="Abstrak Position Subtitle Seperator")
    position_paragraph.add_run(f"{pos_start} - {pos_end}", style="Abstrak Position Subtitle")

    position_paragraph.paragraph_format.space_after = Pt(1)
    # Position Responsibilitites with acomplishments first
    for item in position["accomplishments"]+position["duties"]:
      document.add_paragraph(
        item,
        style='Abstrak Position List Bullet'
      )
    
    last_position = document.paragraphs[-1]
    last_position.paragraph_format.keep_together = True
    #last_position.paragraph_format.space_after = Pt(5)


if __name__ == "__main__":
  # Make sure a data file has been specified.
  # TODO: Maybe use something like Click for input args
  if(len(sys.argv) == 1):
    print("Please specify a data and style file")
    sys.exit(2)

  # Get the resume data
  resume_data = util.load_resume_data(sys.argv[1])
  
  # Create the document object
  document = Document()

  # Configure the styles used within the document
  AbstrakStyle(document, style_path=sys.argv[2])

  # Try build the header
  header_data = resume_data.get("header")
  if header_data is None:
    raise Exception("No header found!")
  header = build_header(document, header_data)

  # Try build the summary
  summary_data = resume_data.get("summary")
  if summary_data is None:
    raise Exception("No summary found!")
  build_summary(document, summary_data)
  
  util.insert_columns_section(document, 2)

  # Try build the highlights
  highlights_data = resume_data.get("highlights")
  if highlights_data is None:
    raise Exception("No highlights found!")
  build_highlights(document, highlights_data)

  util.insert_standard_section(document)

  # Try build the roles
  roles_data = resume_data.get("roles")
  if roles_data is None:
    raise Exception("No roles found!")
  build_roles(document, roles_data)
  
  # Try build the footer
  footer_data = resume_data.get("header")
  if footer_data is None:
    raise Exception("No footer found!")
  build_footer(document, footer_data)

  # Save the resume
  save_file_name = "resume.docx"
  if sys.argv[3]:
    save_file_name = sys.argv[3]

  document.save(save_file_name)